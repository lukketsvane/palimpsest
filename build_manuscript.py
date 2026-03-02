#!/usr/bin/env python3
"""
Build script for PALIMPSEST submission to FormAkademisk journal.
Creates a publication-ready .docx file following journal requirements:
- Single line spacing
- APA 7th edition references
- Anonymized for peer review
- Figures embedded at correct locations
- Block quotes properly formatted
"""

import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
FIGURES_DIR = os.path.join(BASE_DIR, "figures")
OUTPUT_FILE = os.path.join(BASE_DIR, "PALIMPSEST_submission.docx")


def setup_styles(doc):
    """Configure document styles per FormAkademisk requirements."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)
    pf = style.paragraph_format
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)
    pf.line_spacing = Pt(12)  # Single spacing

    # Title style
    if 'ArticleTitle' not in [s.name for s in doc.styles]:
        ts = doc.styles.add_style('ArticleTitle', WD_STYLE_TYPE.PARAGRAPH)
        ts.font.name = 'Times New Roman'
        ts.font.size = Pt(20)
        ts.font.bold = True
        ts.paragraph_format.space_after = Pt(6)
        ts.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Subtitle style
    if 'ArticleSubtitle' not in [s.name for s in doc.styles]:
        ss = doc.styles.add_style('ArticleSubtitle', WD_STYLE_TYPE.PARAGRAPH)
        ss.font.name = 'Times New Roman'
        ss.font.size = Pt(14)
        ss.font.italic = True
        ss.paragraph_format.space_after = Pt(12)

    # Section heading style
    if 'SectionHeading' not in [s.name for s in doc.styles]:
        sh = doc.styles.add_style('SectionHeading', WD_STYLE_TYPE.PARAGRAPH)
        sh.font.name = 'Times New Roman'
        sh.font.size = Pt(12)
        sh.font.bold = True
        sh.paragraph_format.space_before = Pt(18)
        sh.paragraph_format.space_after = Pt(6)
        sh.paragraph_format.keep_with_next = True

    # Subsection heading style
    if 'SubsectionHeading' not in [s.name for s in doc.styles]:
        sub = doc.styles.add_style('SubsectionHeading', WD_STYLE_TYPE.PARAGRAPH)
        sub.font.name = 'Times New Roman'
        sub.font.size = Pt(12)
        sub.font.bold = True
        sub.font.italic = True
        sub.paragraph_format.space_before = Pt(12)
        sub.paragraph_format.space_after = Pt(6)
        sub.paragraph_format.keep_with_next = True

    # Block quote style
    if 'BlockQuote' not in [s.name for s in doc.styles]:
        bq = doc.styles.add_style('BlockQuote', WD_STYLE_TYPE.PARAGRAPH)
        bq.font.name = 'Times New Roman'
        bq.font.size = Pt(10)
        bq.paragraph_format.left_indent = Cm(1.27)
        bq.paragraph_format.space_before = Pt(6)
        bq.paragraph_format.space_after = Pt(6)
        bq.paragraph_format.line_spacing = Pt(10)

    # Figure caption style
    if 'FigureCaption' not in [s.name for s in doc.styles]:
        fc = doc.styles.add_style('FigureCaption', WD_STYLE_TYPE.PARAGRAPH)
        fc.font.name = 'Times New Roman'
        fc.font.size = Pt(10)
        fc.paragraph_format.space_before = Pt(6)
        fc.paragraph_format.space_after = Pt(12)

    # Abstract heading style
    if 'AbstractHeading' not in [s.name for s in doc.styles]:
        ah = doc.styles.add_style('AbstractHeading', WD_STYLE_TYPE.PARAGRAPH)
        ah.font.name = 'Times New Roman'
        ah.font.size = Pt(12)
        ah.font.bold = True
        ah.paragraph_format.space_before = Pt(12)
        ah.paragraph_format.space_after = Pt(6)

    # Keywords style
    if 'Keywords' not in [s.name for s in doc.styles]:
        kw = doc.styles.add_style('Keywords', WD_STYLE_TYPE.PARAGRAPH)
        kw.font.name = 'Times New Roman'
        kw.font.size = Pt(10)
        kw.font.italic = True
        kw.paragraph_format.space_after = Pt(12)

    # Arabic text style
    if 'ArabicText' not in [s.name for s in doc.styles]:
        at = doc.styles.add_style('ArabicText', WD_STYLE_TYPE.PARAGRAPH)
        at.font.name = 'Times New Roman'
        at.font.size = Pt(12)
        at.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        at.paragraph_format.space_before = Pt(6)
        at.paragraph_format.space_after = Pt(3)

    # Poetry translation style
    if 'PoetryTranslation' not in [s.name for s in doc.styles]:
        pt_style = doc.styles.add_style('PoetryTranslation', WD_STYLE_TYPE.PARAGRAPH)
        pt_style.font.name = 'Times New Roman'
        pt_style.font.size = Pt(11)
        pt_style.font.italic = True
        pt_style.paragraph_format.left_indent = Cm(1.27)
        pt_style.paragraph_format.space_after = Pt(6)

    return doc


def add_body_paragraph(doc, text, bold_phrases=None, italic_phrases=None):
    """Add a body paragraph with optional bold/italic phrases."""
    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(3)

    if not bold_phrases and not italic_phrases:
        p.add_run(text)
        return p

    # Simple approach: just add the full text with formatting markers
    # For complex inline formatting, we split on markers
    p.add_run(text)
    return p


def add_body(doc, text, first_indent=True):
    """Add a plain body paragraph."""
    p = doc.add_paragraph(style='Normal')
    if first_indent:
        p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(3)
    p.add_run(text)
    return p


def add_body_with_italic(doc, parts):
    """Add body paragraph with mixed normal/italic runs.
    parts is a list of (text, is_italic) tuples.
    """
    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(3)
    for text, is_italic in parts:
        run = p.add_run(text)
        if is_italic:
            run.italic = True
    return p


def add_block_quote(doc, text):
    """Add a block quote (40+ words, indented, 10pt, no quotation marks)."""
    p = doc.add_paragraph(style='BlockQuote')
    p.add_run(text)
    return p


def add_figure(doc, figure_file, caption_num, caption_text, width=None):
    """Add a figure with caption."""
    filepath = os.path.join(FIGURES_DIR, figure_file)
    if not os.path.exists(filepath):
        print(f"WARNING: Figure not found: {filepath}")
        p = doc.add_paragraph(f"[MISSING FIGURE: {figure_file}]")
        return

    # Add figure
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    if width:
        run.add_picture(filepath, width=width)
    else:
        run.add_picture(filepath, width=Inches(5.5))

    # Add caption
    cap = doc.add_paragraph(style='FigureCaption')
    bold_run = cap.add_run(f"Figur {caption_num}. ")
    bold_run.bold = True
    cap.add_run(caption_text)
    return cap


def add_arabic_and_translation(doc, arabic_text, translation):
    """Add Arabic text (RTL) followed by Norwegian translation."""
    # Arabic
    p = doc.add_paragraph(style='ArabicText')
    run = p.add_run(arabic_text)
    run.font.size = Pt(14)

    # Translation
    t = doc.add_paragraph(style='PoetryTranslation')
    t.add_run(translation)
    return t


def add_section_heading(doc, number, title):
    """Add a numbered section heading."""
    p = doc.add_paragraph(style='SectionHeading')
    p.add_run(f"{number}. {title.upper()}")
    return p


def add_subsection_heading(doc, roman, title):
    """Add a Roman-numeral subsection heading."""
    p = doc.add_paragraph(style='SubsectionHeading')
    p.add_run(f"{roman}. {title}")
    return p


def build_front_matter(doc):
    """Build title page and abstracts."""
    # Title
    p = doc.add_paragraph(style='ArticleTitle')
    p.add_run("PALIMPSEST")

    # Subtitle
    p = doc.add_paragraph(style='ArticleSubtitle')
    p.add_run("Diskurs som overflate, infrastruktur som nerver")

    # Sub-subtitle
    p = doc.add_paragraph(style='Normal')
    run = p.add_run("Ein materialistisk kritikk av Edward Said sitt Orientalism")
    run.italic = True
    p.paragraph_format.space_after = Pt(12)

    # Author (anonymized)
    p = doc.add_paragraph(style='Normal')
    p.add_run("[Forfattar anonymisert for fagfellevurdering]")
    p.paragraph_format.space_after = Pt(6)

    p = doc.add_paragraph(style='Normal')
    p.add_run("[Institusjonstilknyting fjerna]")
    p.paragraph_format.space_after = Pt(18)

    # --- SAMANDRAG ---
    p = doc.add_paragraph(style='AbstractHeading')
    p.add_run("SAMANDRAG")

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.space_after = Pt(6)
    p.add_run(
        "Denne artikkelen utforskar spenninga mellom diskursanalyse og materialistisk kritikk "
        "gjennom eit kunstnarisk prosjekt: ein palimpsest best\u00e5ande av Mahmoud Darwish sitt dikt "
    )
    run = p.add_run("Tibaq")
    run.italic = True
    p.add_run(
        " (2003) inskribert p\u00e5 industriell prismefolie fr\u00e5 ein demontert LCD-skjerm. Med "
        "utgangspunkt i kritikken fr\u00e5 Aijaz Ahmad (1992), Vivek Chibber (2018; 2020) og Michael "
        "Parenti (1995), argumenterer teksten for at Edward Saids "
    )
    run = p.add_run("Orientalism")
    run.italic = True
    p.add_run(
        " (1978) inneheld ein fundamental motseiing mellom eit materialistisk argument \u2014 "
        "orientalismen som konsekvens av kolonialismen \u2014 og eit idealistisk argument \u2014 "
        "orientalismen som \u00e5rsak til kolonialismen. Artikkelen demonstrerer at det andre argumentet "
        "ikkje kan forklare imperialismens historiske m\u00f8nster: statar intervenerer der materielle "
        "interesser er truga, ikkje der orientalistisk diskurs er mest intens. Gjennom ei utvida "
        "lesing av Darwish sin elegi til Said, visar artikkelen korleis poeten implisitt korrigerer "
        "Said sine teoretiske inkonsistensar. Omgrepet \u00aboksidentalisme\u00bb tener som empirisk "
        "motpr\u00f8ve: forvrengde vestlege representasjonar har ikkje produsert oksidentalistisk "
        "imperialisme, noko som indikerer at makt f\u00f8lgjer materielle strukturar, ikkje diskurs. "
        "Verket visualiserer denne kritikken ved \u00e5 la infrastrukturen (prismefolien) fragmentere "
        "diskursen (poesien)."
    )

    # Keywords Norwegian
    p = doc.add_paragraph(style='Keywords')
    bold_run = p.add_run("N\u00f8kkelord: ")
    bold_run.bold = True
    bold_run.italic = False
    p.add_run("orientalisme, materialisme, imperialisme, palimpsest, Edward Said, Mahmoud Darwish, kontrapunkt")

    # --- ABSTRACT ---
    p = doc.add_paragraph(style='AbstractHeading')
    p.add_run("ABSTRACT")

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.space_after = Pt(6)
    p.add_run(
        "This article explores the tension between discourse analysis and materialist critique through "
        "an artistic project: a palimpsest consisting of Mahmoud Darwish\u2019s poem "
    )
    run = p.add_run("Tibaq")
    run.italic = True
    p.add_run(
        " (2003) inscribed on industrial prism film from a dismantled LCD screen. Drawing on the "
        "critiques of Aijaz Ahmad (1992), Vivek Chibber (2018; 2020) and Michael Parenti (1995), the "
        "text argues that Edward Said\u2019s "
    )
    run = p.add_run("Orientalism")
    run.italic = True
    p.add_run(
        " (1978) contains a fundamental contradiction between a materialist argument \u2014 Orientalism "
        "as a consequence of colonialism \u2014 and an idealist argument \u2014 Orientalism as a cause "
        "of colonialism. The article demonstrates that the latter cannot explain historical patterns "
        "of imperialism. Through an extended reading of Darwish\u2019s elegy for Said, the article shows "
        "how the poet implicitly corrects Said\u2019s theoretical inconsistencies. The concept of "
        "\u00abOccidentalism\u00bb serves as an empirical counter-test: distorted Western representations "
        "have not produced Occidentalist imperialism, indicating that power follows material structures, "
        "not discourse."
    )

    # Keywords English
    p = doc.add_paragraph(style='Keywords')
    bold_run = p.add_run("Keywords: ")
    bold_run.bold = True
    bold_run.italic = False
    p.add_run("Orientalism, materialism, imperialism, palimpsest, Edward Said, Mahmoud Darwish, counterpoint")

    return doc


def build_section_1(doc):
    """1. INNLEIING: EIT PALIMPSEST SOM ARGUMENT"""
    add_section_heading(doc, 1, "Innleiing: Eit palimpsest som argument")

    p = add_body(doc,
        "Verket som f\u00f8lgjer denne artikkelen er ein materiell palimpsest, men ikkje av pergament. "
        "Overflata er ei optisk prismeplate, henta ut av ein demontert LCD-skjerm, sj\u00f8lve den "
        "fysiske komponenten som har som oppg\u00e5ve \u00e5 spreie og homogenisere lyset i v\u00e5re "
        "digitale representasjonsmaskiner. P\u00e5 denne industrielle plastfolien, ripete og merka av "
        "bruk, er Mahmoud Darwish sitt dikt ", first_indent=False)
    # We need to add italic "Tibaq" inline - let's rebuild this paragraph
    doc.paragraphs[-1]._element.getparent().remove(doc.paragraphs[-1]._element)

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.space_after = Pt(3)
    p.add_run(
        "Verket som f\u00f8lgjer denne artikkelen er ein materiell palimpsest, men ikkje av pergament. "
        "Overflata er ei optisk prismeplate, henta ut av ein demontert LCD-skjerm, sj\u00f8lve den "
        "fysiske komponenten som har som oppg\u00e5ve \u00e5 spreie og homogenisere lyset i v\u00e5re "
        "digitale representasjonsmaskiner. P\u00e5 denne industrielle plastfolien, ripete og merka av "
        "bruk, er Mahmoud Darwish sitt dikt "
    )
    run = p.add_run("Tibaq")
    run.italic = True
    p.add_run(
        " (2003) skrive for hand med svart tusj. Skrifta flyt over dei mikroskopiske rillene som "
        "vanlegvis styrer lyset mot auga v\u00e5re. Her er det ingen simulert \u00abkald cyan\u00bb "
        "infrastruktur som ligg under; infrastrukturen er sj\u00f8lve mediet. Det er s\u00f8ppel fr\u00e5 "
        "informasjonssamfunnet, gjenbrukt som b\u00e6raranordning for poesi."
    )

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(3)
    p.add_run(
        "Oppg\u00e5va sp\u00f8r etter eit \u00aboksidentaliserande\u00bb uttrykk, korleis det ville "
        "sett ut om rollene var reverserte, om \u00ab\u00d8sten\u00bb konstruerte eit forenkla bilete "
        "av \u00abVesten\u00bb. Men for \u00e5 svare p\u00e5 dette sp\u00f8rsm\u00e5let m\u00e5 me "
        "fyrst unders\u00f8kje premissen det kviler p\u00e5: Edward Said sin teori om orientalisme, og "
        "dei grunnleggjande problema med denne teorien som eit rammeverk for \u00e5 forst\u00e5 makt "
        "og dominans."
    )

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(3)
    p.add_run(
        "Dette notatet argumenterer for at Said sitt verk, trass sin enorme innverknad, kviler p\u00e5 "
        "ein djup indre motseiing som undergrev dets eige kritiske potensial. Ved \u00e5 lokalisere makt "
        "i diskurs og representasjon heller enn i materielle strukturar, endar Said opp med \u00e5 tilby "
        "ein kulturkritikk som fungerer som avleiingsman\u00f8ver fr\u00e5 dei faktiske mekanismane for "
        "imperialisme. Verket mitt fors\u00f8kjer \u00e5 visualisere denne kritikken gjennom palimpsesten "
        "si doble lesing: det vakre kulturelle laget og det brutale infrastrukturelle laget, eksisterande "
        "side om side utan \u00e5 p\u00e5verke kvarandre."
    )

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(3)
    p.add_run("Denne analysen byggjer p\u00e5 ein rik tradisjon av materialistisk kritikk, fr\u00e5 Aijaz Ahmad sin omfattande gjennomgang i ")
    run = p.add_run("In Theory")
    run.italic = True
    p.add_run(" (1992) til Vivek Chibber sin nyare \u00abOrientalism and Its Afterlives\u00bb (2020). Men den hentar ogs\u00e5 innsikter fr\u00e5 Michael Parenti sin ")
    run = p.add_run("Against Empire")
    run.italic = True
    p.add_run(" (1995), eit verk som tilbyr ein presis definisjon av kva imperialisme ")
    run = p.add_run("faktisk er")
    run.italic = True
    p.add_run(", og dermed eit kriterium for \u00e5 vurdere om Said sitt rammeverk fangar dei vesentlege eigenskapane ved fenomenet det p\u00e5st\u00e5r \u00e5 analysere.")

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(3)
    p.add_run("Samstundes dreg denne analysen p\u00e5 nyare spekulativ filosofi, s\u00e6rleg Armen Avanessian og Andreas T\u00f6pfer (2014) sitt arbeid med ")
    run = p.add_run("Speculative Drawing")
    run.italic = True
    p.add_run(" og den breie accelerationist-tradisjonen dokumentert i ")
    run = p.add_run("#Accelerate: The Accelerationist Reader")
    run.italic = True
    p.add_run(" (Mackay & Avanessian, 2014). Avanessian og Anke Hennig (2015) si utforsking av tid og poetikk i ")
    run = p.add_run("Present Tense: A Poetics")
    run.italic = True
    p.add_run(" tilbyr teoretiske reiskapar for \u00e5 forst\u00e5 korleis kritisk praksis kan produsere, ikkje berre representere, verkelegheit.")

    # Figure 2: Representasjonsmaskinene (figur_01.png = 5 heads reading different media)
    add_figure(doc, "figur_01.png", 2,
        "Representasjonsmaskinene si utvikling: fr\u00e5 bok til smarttelefon. Mediet endrar seg, men den strukturelle relasjonen mellom betraktar og representasjon forblir den same. Kjelde: Eige illustrasjon.",
        width=Inches(5.5))


def build_section_2(doc):
    """2. VERKET: MATERIALITET OG KONSTRUKSJON"""
    add_section_heading(doc, 2, "Verket: Materialitet og konstruksjon")

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.space_after = Pt(3)
    p.add_run("Palimpsesten er konstruert av tre hovudlag (sj\u00e5 Figur 4). Det nedste laget er ei MDF-bakplate som gjev strukturell stabilitet. Over denne ligg papirlaget med Darwish sin tekst, handskriven i svart tusj. Det \u00f8vste laget er prismefolien, henta fr\u00e5 ein demontert LCD-skjerm \u2014 den optiske komponenten som i sin opphavlege funksjon spreier og homogeniserer bakgrunnsbelysning til ein jamn lysflate.")

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(3)
    p.add_run("Prismefolien er ikkje eit n\u00f8ytralt medium. Dei mikroskopiske rillene i overflata bryt lyset i spektralfargar n\u00e5r det passerer gjennom (sj\u00e5 Figur 3). Denne optiske eigenskapen, designa for industriell effektivitet, blir i verket ein visuell metafor for forholdet mellom diskurs og infrastruktur: kalligrafien (diskursen) vert fragmentert og forvrengd av dei materielle strukturane han er inskribert p\u00e5. Materialet ber sine eigne spor \u2014 folien er ripete og sliten fr\u00e5 si tidlegare tenestegjering inne i ein skjerm, merke av bruk som ingen var meint \u00e5 sj\u00e5. Desse ripene er like mykje ein del av verket som den kalligraferte teksten.")

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(3)
    p.add_run("Valet av materiale er sj\u00f8lve argumentet: i staden for \u00e5 skrive p\u00e5 pergament eller papir i ei orientaliserande imitering av islamsk bokkunst, er teksten lagt p\u00e5 det industrielle avfallet fr\u00e5 vestleg bildeproduksjon. Verket nektar nostalgi og insisterer p\u00e5 det samtidige. Det er s\u00f8ppel gjenbrukt som b\u00e6rar for poesi \u2014 og nettopp denne spenninga mellom det poetiske innhaldet og det infrastrukturelle mediet utgjer palimpsesten si doble lesing.")

    # Figure 3: Den optiske effekten (figur_02.png = eye/prism diagram)
    add_figure(doc, "figur_02.png", 3,
        "Den optiske effekten som teoretisk modell. Utan prismefolie: uforstyrra kulturell representasjon (idealistisk diskurs). Gjennom prismefolien: diskursen fragmentert av materielle strukturar (materialistisk r\u00f8yndom). Det analytiske blikket ser gjennom infrastrukturen, men infrastrukturen endrar det ein ser. Kjelde: Eige illustrasjon.",
        width=Inches(5.5))

    # Figure 4: Samansetjinga (figur_03.png = exploded view with 3 layers)
    add_figure(doc, "figur_03.png", 4,
        "Samansetjinga av palimpsesten i eksplodert visning. Dei tre laga \u2014 prismefolie (fr\u00e5 LCD-skjerm), kalligrafi p\u00e5 papir, og MDF-bakplate \u2014 utgjer ein materiell heilskap der kvart lag har sin eigen funksjon og symbolikk. Kjelde: Eige illustrasjon.",
        width=Inches(5))


def build_section_3(doc):
    """3. SAID SIN DOBLE TESE: EIN INDRE MOTSEIING"""
    add_section_heading(doc, 3, "Said sin doble tese: Ein indre motseiing")

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.space_after = Pt(3)
    p.add_run("Edward Said sitt ")
    run = p.add_run("Orientalism")
    run.italic = True
    p.add_run(" (1978) har hatt ein n\u00e6rast uovertruffen innverknad p\u00e5 humaniora og samfunnsvitskap dei siste femti \u00e5ra. Boka argumenterer for at Vesten har konstruert eit systematisk forvrengt bilete av \u00abOrienten\u00bb som eksotisk, irrasjonell, feminin og tilbakest\u00e5ande, eit bilete som har tent til \u00e5 rettferdiggjere kolonial dominans. Said (1978, s. 5) skriv i innleiinga:")

    add_block_quote(doc,
        "The relationship between Occident and Orient is a relationship of power, of domination, of varying degrees of a complex hegemony.")

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(3)
    p.add_run("Men som Vivek Chibber har vist i sin analyse av verkets doble arv, inneheld ")
    run = p.add_run("Orientalism")
    run.italic = True
    p.add_run(" to fundamentalt ulike argument som st\u00e5r i spenning med kvarandre (Chibber, 2020). \u00c5 forst\u00e5 denne spenninga er avgjerande for \u00e5 vurdere verkets teoretiske haldbarheit.")

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(3)
    p.add_run("Det fyrste argumentet er materialistisk i sin grunnstruktur: orientalistisk diskurs oppstod som ein konsekvens av kolonialismen. Kolonimaktene trong ideologisk rettferdiggjering for si dominans, og orientalismen leverte denne rettferdiggjeringa. Her er forholdet mellom makt og kunnskap klart: materielle interesser skapar sin legitimerande ideologi. Chibber (2020) forklarer:")

    add_block_quote(doc,
        "Said\u2019s argument here is a fairly traditional, materialist explanation for how and why Orientalist ideology came to occupy such a prominent place in European culture in the modern period. Just as any system of domination creates an ideological discourse to justify and naturalize its superordinate position, so too colonialism created a legitimizing discourse of its own.")

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(3)
    p.add_run("Dette argumentet er, som Chibber p\u00e5peikar, konvensjonelt og velkjent. Det plasserer Said trygt innanfor den marxistiske tradisjonen der ideologi forst\u00e5st som overbygning over ein materiell basis.")

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(3)
    p.add_run("Det andre argumentet inverterer dette forholdet fullstendig. Her blir orientalismen ikkje berre ein konsekvens av kolonialismen, men ein \u00e5rsak til han. Said introduserer omgrepet \u00ablatent orientalisme\u00bb for \u00e5 beskrive ei djuptliggande vestleg haldning til \u00d8sten som eksisterer n\u00e6rast uavhengig av spesifikke historiske omstende. Denne latente orientalismen, hevdar Said, er s\u00e5 djupt forankra i vestleg kultur, fr\u00e5 Aiskhylos til Nixon, at den i seg sj\u00f8lv genererer kolonial aggresjon. Chibber (2020, med sitat fr\u00e5 Said, 1978, s. 207) siterer Said:")

    add_block_quote(doc,
        "Latent Orientalism came packaged as a \u2018will to power\u2019 \u2026 Hence the obsessive accumulation of facts, Said suggests, \u2018made Orientalism fatally tend towards the systematic accumulation of human beings and territories.\u2019")

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(3)
    p.add_run("Her er forholdet mellom diskurs og makt invertert: i staden for at eit system av dominans skapar sin legitimerande ideologi, er det ideologien som genererer dominansen. Said (1978) formulerer det slik:")

    add_block_quote(doc,
        "Orientalism is not really a thoroughly modern phenomenon \u2026 but is the natural product of an ancient and almost irresistible European bent of mind to misrepresent the realities of other cultures.")

    # Figure 5: Said sin metode
    add_figure(doc, "figur_04.png", 5,
        "Said sin metode (venstre): tekstar og \u00f8konomiske fakta som separate domene. Den materialistiske kritikken (h\u00f8gre): korleis nonsens vert transformert til tesar gjennom akademia, legitimert av imperialistisk vald. Kjelde: Eige illustrasjon.",
        width=Inches(5))


def build_section_4(doc):
    """4. DEN TIDLEGE KRITIKKEN: AL-AZM OG AHMAD"""
    add_section_heading(doc, 4, "Den tidlege kritikken: al-Azm og Ahmad")

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.space_after = Pt(3)
    p.add_run("Sadik Jalal al-Azm identifiserte denne motseiinga allereie i 1981, berre tre \u00e5r etter at ")
    run = p.add_run("Orientalism")
    run.italic = True
    p.add_run(" vart publisert. I sin skarpe kritikk i tidsskriftet ")
    run = p.add_run("Khamsin")
    run.italic = True
    p.add_run(" p\u00e5peikte al-Azm at om orientalismen er ein djuptliggande del av det vestlege kognitive apparatet, slik Said sitt andre argument hevdar, korleis kan d\u00e5 nokon vestleg forfattar, inkludert Said sj\u00f8lv, nokon gong unnslippe han? Al-Azm (2000, s. 220) observerte at Said si eiga formulering \u2014 at orientalismen er \u00abthe natural product of an ancient and almost irresistible European bent of mind to misrepresent the realities of other cultures\u00bb (Said, 1978, s. 204) \u2014 f\u00f8rte til ein sj\u00f8lvmotseiande konklusjon.")

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(3)
    p.add_run("Aijaz Ahmad tok opp same tr\u00e5den i sin omfattande kritikk fr\u00e5 1992. I kapittelet \u00abOrientalism and After\u00bb i ")
    run = p.add_run("In Theory")
    run.italic = True
    p.add_run(" spekulerte Ahmad i at Said sitt andre argument kanskje kunne tilskrivast innverknaden fr\u00e5 Michel Foucault, men han stilte sp\u00f8rsm\u00e5l ved om Foucault sj\u00f8lv ville st\u00f8tta ideen om ein p\u00e5st\u00e5tt kontinuitet i vestleg diskurs fr\u00e5 Homer til Richard Nixon. Ahmad (2000, s. 285) skreiv:")

    add_block_quote(doc,
        "Said\u2019s focus on text \u2018facilitates a reading of history not from the basis of material production, but from its systems of representations.\u2019")

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(3)
    p.add_run("Dette er eit avgjerande punkt. Ahmad identifiserer det sentrale problemet med Said sitt prosjekt: ved \u00e5 fokusere p\u00e5 tekst og representasjon, mister ein dei materielle produksjonsforholda av syne. Analysen blir ahistorisk fordi den ikkje kan forklare kvifor orientalismen oppstod p\u00e5 eit spesifikt historisk tidspunkt, under spesifikke materielle vilk\u00e5r.")

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(3)
    p.add_run("Ahmad (1992, s. 184) reiste to sentrale innvendingar mot Said sitt andre argument. For det fyrste syntest Said \u00e5 ta det orientalistiske tankem\u00f8nsteret for \u00e5 vera s\u00e5 gjennomgripande i omfang og s\u00e5 mektig i innverknad at moglegheita for \u00e5 unnslippe det verka ekstremt fjern. For det andre, ved \u00e5 forkaste den tradisjonelle marxistiske vektlegginga av kolonialismen sin materielle basis, gjorde Said det umogleg \u00e5 identifisere dei konkrete interessene som driv imperialismen.")

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(3)
    p.add_run("Ahmad (1992, s. 221) p\u00e5peikte ogs\u00e5 at Said sine selektive lesingar f\u00f8rte til at sj\u00f8lv radikale kritikarar av imperialismen vart assimilerte inn i den orientalistiske tradisjonen. Marx, til d\u00f8mes, blir av Said framstilt som ein orientalist, trass i at Marx sin kritikk av det britiske imperiet i India var grunnleggjande materialistisk og ikkje kvilte p\u00e5 nokon p\u00e5stand om \u00aborientalsk\u00bb essens.")


def build_section_5(doc):
    """5. KVA IMPERIALISME FAKTISK ER: PARENTI SIN DEFINISJON"""
    add_section_heading(doc, 5, "Kva imperialisme faktisk er: Parenti sin definisjon")

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.space_after = Pt(3)
    p.add_run("For \u00e5 forst\u00e5 kvifor Said sitt andre argument er s\u00e5 problematisk, m\u00e5 me fyrst klargjere kva imperialisme faktisk inneber. Michael Parenti (1995, s. 1) tilbyr ein presis definisjon i opninga av ")
    run = p.add_run("Against Empire")
    run.italic = True
    p.add_run(":")

    add_block_quote(doc,
        "By \u2018imperialism\u2019 I mean the process whereby the dominant politico-economic interests of one nation expropriate for their own enrichment the land, labor, raw materials, and markets of another people.")

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(3)
    p.add_run("Denne definisjonen er avgjerande fordi han lokaliserer imperialismen i materielle prosessar, ikkje i representasjonar eller diskursar. Imperialisme handlar om ekspropriasjon av land, utbytting av arbeidskraft, utvinning av r\u00e5varer og kontroll over marknader. Desse prosessane kan sj\u00f8lvsagt legitimerast gjennom ideologiar, inkludert orientalistiske, men ideologiane er ikkje det som driv dei.")

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(3)
    p.add_run("Parenti (1995, s. 15) gjer eit vesentleg poeng om forholdet mellom kapitalisme og imperialisme som direkte utfordrar Said sitt rammeverk:")

    add_block_quote(doc,
        "Whether imperialism is necessary for capitalism is really not the question. Many things that are not absolutely necessary are still highly desirable, therefore strongly preferred and vigorously pursued.")

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(3)
    p.add_run("Det er ikkje slik at kapitalistar fyrst trur orientalistiske ting om andre folk, og deretter utbyttar dei. Det er omvendt: dei utbyttar fordi det er profitabelt, og orientalismen fungerer som ideologisk st\u00f8nad for denne utbyttinga. Parenti (1995, s. 7) dokumenterer korleis denne dynamikken har fungert historisk:")

    add_block_quote(doc,
        "By 1850, India\u2019s debt had grown to \u00a353 million. From 1850 to 1900, its per capita income dropped by almost two-thirds. The massive poverty we associate with India was not that country\u2019s original historical condition. British imperialism did two things: first, it ended India\u2019s development, then it forcibly underdeveloped that country.")

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(3)
    p.add_run("Desse tala, denne konkrete dokumentasjonen av utbytting, er det som forklarer imperialismen. Ikkje tekstar om India sin p\u00e5st\u00e5tte \u00aborientalske\u00bb karakter, men gjeldsmekanismar, tvangseksport og systematisk underutvikling.")

    # Figure 6: Den koloniale relasjonen
    add_figure(doc, "figur_05.png", 6,
        "Den koloniale relasjonen: imperiet som subjekt, kolonien som objekt. Diskursen skjuler r\u00f8ynda: \u00abvi studerer maska, dei tek oljen.\u00bb Kjelde: Eige illustrasjon.",
        width=Inches(4.5))


def build_section_6(doc):
    """6. DEN EMPIRISKE TESTEN: KVAR INTERVENERER IMPERIET?"""
    add_section_heading(doc, 6, "Den empiriske testen: Kvar intervenerer imperiet?")

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.space_after = Pt(3)
    p.add_run("Parenti (1995, s. 43) tilbyr ein enkel men avgjerande empirisk test for \u00e5 vurdere kva som faktisk driv imperialistisk intervensjon:")

    add_block_quote(doc,
        "Interventionist forces do not go where capital exists as such; they go where capital is threatened. They have not intervened in affluent Switzerland, for instance, because capitalism in that country is relatively secure and unchallenged.")

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(3)
    p.add_run("Om Said sitt andre argument er korrekt, burde orienteringa av intervensjonar f\u00f8lgje intensiteten i orientalistiske representasjonar. Land som er sterkt \u00aborientaliserte\u00bb i vestleg diskurs burde vera meir utsette for intervensjon enn land som ikkje er det. Men det er ikkje det me observerer.")

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(3)
    p.add_run("USA intervenerer der det finst strategiske ressursar, der det finst trugsmål mot kapitalinteresser, der det finst geopolitiske m\u00e5l. Dei intervenerer ikkje i land som manglar slike interesser, uansett kor \u00aborientalistisk\u00bb diskursen om desse landa m\u00e5tte vera. Omvendt intervenerer dei villig i land som knapt figurerer i orientalistisk diskurs, s\u00e5 lenge dei materielle interessene er til stades.")

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(3)
    p.add_run("Parenti (1995, s. 9) dokumenterer dette m\u00f8nsteret. Dei faktiske mekanismane for imperialisme \u2014 gjeldsmekanismar, handelsavtalar, sanksjonar, milit\u00e6rbasar, etterretningsoperasjonar, st\u00f8tte til kompradorklassar \u2014 krev ikkje orientalistiske representasjonar for \u00e5 fungere. Dei opererer like effektivt i Latin-Amerika som i Midt\u00f8sten, like effektivt mot kristne som mot muslimar, like effektivt mot \u00abVesten\u00bb sine eigne arbeidsklassar som mot folk i det globale s\u00f8r.")

    # Figure 7: Orientalisme som diskurs
    add_figure(doc, "figur_06.png", 7,
        "Orientalisme som diskurs, fr\u00e5 simulering til imitasjon: den skjulte r\u00f8yndomen (basis) er p\u00e5tatt. Kjelde: Eige illustrasjon.",
        width=Inches(4.5))


def build_section_7(doc):
    """7. CHOMSKY MOT FOUCAULT: EMPIRISME MOT DISKURSTEORI"""
    add_section_heading(doc, 7, "Chomsky mot Foucault: Empirisme mot diskursteori")

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.space_after = Pt(3)
    p.add_run("Den teoretiske konflikten mellom materialistisk analyse og diskursteori vart dramatisk iscenesett i den ber\u00f8mte debatten mellom Noam Chomsky og Michel Foucault i 1971. Said var, etter eige utsegn, p\u00e5verka av \u00abto polaritetar\u00bb, Foucault p\u00e5 den eine sida og Chomsky p\u00e5 den andre (Ahmad, 1992, s. 165).")

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(3)
    p.add_run("Chomsky sitt standpunkt var klart: makt m\u00e5 analyserast empirisk, gjennom dokumentasjon av faktiske handlingar, institusjonar og interesser. Foucault, derimot, insisterte p\u00e5 at makt opererer gjennom diskursive formasjonar. I debatten uttrykte Foucault ein posisjon som n\u00e6rast kollapsar i nihilisme n\u00e5r det gjeld sp\u00f8rsm\u00e5let om rettferd (Chomsky & Foucault, 2006).")

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(3)
    p.add_run("Chomsky si innvending var skarp. I eit seinare intervju oppsummerte han problemet med den poststrukturalistiske tiln\u00e6rminga. Chomsky (2011) skreiv:")

    add_block_quote(doc,
        "When [Foucault\u2019s ideas are] decoded, they say nothing. Strip away the jargon: \u2018power circulates,\u2019 \u2018knowledge is implicated in domination.\u2019 Either truisms or mystification.")

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(3)
    p.add_run("Problemet er ikkje at Foucault tek feil i at makt opererer gjennom diskurs. Makt gjer det. Problemet er at ved \u00e5 gjere diskurs til det sentrale analyseobjektet, risikerer ein \u00e5 miste dei materielle strukturane av syne. Ein kan analysere orientalistiske tekstar i det uendelege utan \u00e5 nokon gong konfrontere Pentagon, IMF, Verdsbanken eller oljeselskapa.")

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("Orientalism")
    run.italic = True
    p.add_run(", s\u00e6rleg i sitt andre argument, heller tydeleg mot Foucault. Fred Halliday observerte, som sitert i Chibber (2020):")

    add_block_quote(doc,
        "While much of the other work was framed in broadly Marxist terms and was a universalist critique, Said, eschewing materialist analysis, sought to apply literary critical methodology and to offer an analysis specific to something called \u2018the Orient.\u2019")

    # Figure 8: Den koloniale relasjonen
    add_figure(doc, "figur_07.png", 8,
        "Den koloniale relasjonen: oksidenten (subjekt) og orienten (objekt). Representasjonen kan tolkast som bedrevitande eller sk\u00e5nsam, men den strukturelle asymmetrien forblir den same. Kjelde: Eige illustrasjon.",
        width=Inches(4.5))


def build_section_8(doc):
    """8. CHIBBER SIN ANALYSE: DEN DOBLE ARVEN"""
    add_section_heading(doc, 8, "Chibber sin analyse: Den doble arven")

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.space_after = Pt(3)
    p.add_run("Vivek Chibber sin artikkel \u00abOrientalism and Its Afterlives\u00bb (2020) tilbyr kanskje den mest presise analysen av spenninga i Said sitt verk. Om \u00ablatent orientalisme\u00bb har eksistert i vestleg kultur sidan Aiskhylos, kvifor venta d\u00e5 europeisk kolonialisme til det syttande og attande hundre\u00e5ret med \u00e5 ekspandere globalt?")

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(3)
    p.add_run("Chibber (2020) sin analyse presiserer:")

    add_block_quote(doc,
        "If Said were to agree that, even if Orientalism had not been available as an academic discipline, even if latent Orientalism had been absent from the scene, its basic elements could have nonetheless been crafted ex nihilo in order to justify colonial rule, then he would be suggesting that latent Orientalism was not in fact a necessary part of the causal complex that brought about colonialism.")

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(3)
    p.add_run("Chibber (2018) p\u00e5peikar ogs\u00e5 den problematiske essensialismen som f\u00f8lgjer av Said sitt andre argument:")

    add_block_quote(doc,
        "If, as Said and others admit, all cultures express wayward and at times racist ideas of the \u2018other,\u2019 one needs extra-cultural explanations to uncover why colonialism and imperialism in the modern age were undertaken by European powers.")

    # Figure 9: Det orientalistiske prismet
    add_figure(doc, "figur_08.png", 9,
        "Det orientalistiske prismet: korleis imperielle interesser konstruerer \u00abOrienten\u00bb som skjerm mellom betraktaren og den materielle r\u00f8yndomen. Kjelde: Eige illustrasjon.",
        width=Inches(4.5))


def build_section_9(doc):
    """9. DEI MATERIELLE KOSTNADANE: SMITH SIN DOKUMENTASJON"""
    add_section_heading(doc, 9, "Dei materielle kostnadane: Smith sin dokumentasjon")

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.space_after = Pt(3)
    p.add_run("David Michael Smith (2023, s. 8) har i ")
    run = p.add_run("Endless Holocausts")
    run.italic = True
    p.add_run(" dokumentert omfanget av imperialismen i reine menneskelege kostnadar:")

    add_block_quote(doc,
        "The US Empire is maintained by a network of client states that \u2018encompass 40 per cent of the world\u2019s countries.\u2019 It is supported by 800 foreign bases, with 200,000 military staff and contractors deployed in 140 countries.")

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(3)
    p.add_run("Desse tala representerer ikkje representasjonar. Dei representerer faktiske d\u00f8dsfall, faktiske basar, faktiske troppar, faktisk infrastruktur for dominans. Parenti (2010) kallar imperialismen sin faktiske karakter:")

    add_block_quote(doc,
        "Empires impoverish whole populations and kill lots and lots of innocent people. This is another thing that empires do which too often goes unmentioned in the historical and political literature.")

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(3)
    p.add_run("Desse d\u00f8dsfalla skjedde ikkje p\u00e5 grunn av orientalistiske representasjonar. Dei skjedde p\u00e5 grunn av materiell ekspansjon, ressursutvinning og geopolitisk kontroll.")


def build_section_10(doc):
    """10. SAMTIDSIMPERIALISME: FR\u00c5 TEORI TIL R\u00d8YND"""
    add_section_heading(doc, 10, "Samtidsimperialisme: Fr\u00e5 teori til r\u00f8ynd")

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.space_after = Pt(3)
    p.add_run("Medan akademikarar debatterer representasjonar, utfaldar imperialismen seg i sanntid. I januar 2026 kidnappa det amerikanske milit\u00e6ret Venezuelas president Nicol\u00e1s Maduro i ein operasjon Trump sj\u00f8lv skildra som ut\u00f8ving av det han kallar \u00abDonroe Doctrine\u00bb (ABC News, 2026). Samstundes held Trump fram med \u00e5 true Danmark med anneksjon av Gr\u00f8nland og opne konfrontasjonar med NATO-allierte.")

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(3)
    p.add_run("Desse hendingane illustrerer det materialistiske argumentet med brutal klarheit. Trump snakkar ikkje om kulturelle representasjonar av Venezuela eller Gr\u00f8nland. Han snakkar om olje, mineral og geopolitisk kontroll. Som han sa om Venezuela: \u00abThey have the largest oil reserves in the world\u00bb (ABC News, 2026). Som Parenti (1995) formulerte det: imperiet forf\u00f8lgjer ikkje \u00abpower for power\u2019s sake\u00bb, men reelle og enorme materielle interesser.")

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(3)
    p.add_run("I Gaza utfaldar den same logikken seg. Trump har f\u00f8resl\u00e5tt \u00e5 ta over Gazastripa, tvangsflytte den palestinske befolkninga, og gjere omr\u00e5det om til ei \u00abspesiell \u00f8konomisk sone\u00bb. Her blir orientalistisk diskurs om palestinarar sekund\u00e6r til dei konkrete planane om territorium og kontroll.")

    # Figure 10 - IMAGE FILE MISSING (no matching PNG found in figures/)
    # TODO: Add figur image file for Figur 10
    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("[Figur 10: Bilete manglar]")
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(150, 150, 150)
    cap = doc.add_paragraph(style='FigureCaption')
    cap.add_run("Figur 10. Det finst ingen transcendent dimensjon i kapitalismen (venstre): figuren dabbar i tomrommet. Den historiske prosessen har ingenting å gjere med gjennomføringa av eit ideal (høgre). Kjelde: Eige illustrasjon.")

def build_section_11(doc):
    """11. OKSIDENTALISME: \u00c5 SNU SAID P\u00c5 HOVUDET"""
    add_section_heading(doc, 11, "Oksidentalisme: \u00c5 snu Said p\u00e5 hovudet")

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.space_after = Pt(3)
    p.add_run("No kan me vende tilbake til oppg\u00e5va sitt sp\u00f8rsm\u00e5l: korleis ville eit \u00aboksidentaliserande\u00bb uttrykk sj\u00e5 ut? Ian Buruma og Avishai Margalit (2004) har i ")
    run = p.add_run("Occidentalism")
    run.italic = True
    p.add_run(" dokumentert eksakt slike representasjonar. \u00abVesten\u00bb blir i visse ikkje-vestlege diskursar framstilt som kald, sjellaus, materialistisk, \u00e5ndlaus, mekanisk, dekadent.")

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(3)
    p.add_run("Her kjem den avgjerande innsikta fr\u00e5 den materialistiske kritikken: det spelar inga rolle. Om Said sitt rammeverk er korrekt, burde oksidentalisme vera like problematisk som orientalisme. Begge er forvrengde representasjonar som, if\u00f8lgje Said sin logikk, skulle produsere maktrelasjonar. Men kvar er det oksidentalistiske imperiet?")

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(3)
    p.add_run("Svaret er at makt ikkje f\u00f8lgjer representasjon. Makt f\u00f8lgjer materielle forhold: \u00f8konomisk kapasitet, milit\u00e6r styrke, tilgang til ressursar, kontroll over finansielle institusjonar. Oksidentalistiske representasjonar eksisterer i rikeleg mon, men dei har ikkje produsert nokon oksidentalistisk imperialisme fordi dei manglande samfunna har mangla den materielle kapasiteten til \u00e5 omsette representasjonar i dominans.")

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(3)
    p.add_run("Om Chomsky, Ahmad og Parenti sin kritikk er korrekt, er heile rammeverket av representasjonskritikk ein avleiingsman\u00f8ver. Oljeleidningane g\u00e5r same vegen uansett kva folk trur om kvarandre. Finansstraumane f\u00f8lgjer sin eigen logikk. Milit\u00e6rbasane ligg der dei ligg p\u00e5 grunn av strategiske kalkulasjonar, ikkje p\u00e5 grunn av kulturelle representasjonar.")


def build_section_12(doc):
    """12. DARWISH LES SAID MOT SAID: TIBAQ SOM MATERIALISTISK KORREKTIV"""
    add_section_heading(doc, 12, "Darwish les Said mot Said: Tibaq som materialistisk korrektiv")

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.space_after = Pt(3)
    p.add_run("Det er ei underleg ironi i at det kanskje mest presise svaret p\u00e5 den indre motseiinga i Said sitt ")
    run = p.add_run("Orientalism")
    run.italic = True
    p.add_run(" ikkje kjem fr\u00e5 ein akademikar, men fr\u00e5 ein poet. Mahmoud Darwish sitt ")
    run = p.add_run("Tibaq")
    run.italic = True
    p.add_run(" (\u0637\u0628\u0627\u0642, 2003/2005), skrive som ein elegi til Said kort tid etter hans d\u00f8d, fungerer samstundes som hyllest og som implisitt kritikk.")

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(3)
    p.add_run("Avanessian og T\u00f6pfer (2014) argumenterer i ")
    run = p.add_run("Speculative Drawing")
    run.italic = True
    p.add_run(" for at teikning og poesi ikkje illustrerer teori; dei ")
    run = p.add_run("er")
    run.italic = True
    p.add_run(" ei form for tenking. Teikningane \u00abprovide an occasion to think about thinking \u2014 a speculative thinking and writing in concept and through images\u00bb (s. 12). P\u00e5 same vis fungerer Darwish sitt dikt ikkje som illustrasjon til Said sin teori, men som ein autonom intellektuell praksis som tenkjer ")
    run = p.add_run("med")
    run.italic = True
    p.add_run(" Said samstundes som det tenkjer ")
    run = p.add_run("mot")
    run.italic = True
    p.add_run(" han.")

    # Definition box for Kontrapunkt
    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(1.27)
    run = p.add_run("Kontrapunkt")
    run.bold = True
    p.add_run(" (")
    run = p.add_run("tibaq")
    run.italic = True
    p.add_run(", \u0637\u0628\u0627\u0642): Retorisk og musikalsk omgrep for \u00e5 halde fleire stemmer eller narrativ i spenning samstundes, utan \u00e5 l\u00f8yse dei opp i syntese.")

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(3)
    p.add_run("Said l\u00e5nte omgrepet fr\u00e5 musikkteori i ")
    run = p.add_run("Culture and Imperialism")
    run.italic = True
    p.add_run(" (1993) for \u00e5 beskrive ein lesem\u00e5te der kolonisator og kolonisert, sentrum og periferi, blir haldne i hovudet samstundes.")

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(3)
    p.add_run("Najat Rahman (2007) sin analyse i ")
    run = p.add_run("PMLA")
    run.italic = True
    p.add_run(" visar korleis diktet representerer ei radikal omforming av den arabiske elegien (")
    run = p.add_run("marthiya")
    run.italic = True
    p.add_run("). Men det er meir enn ei generisk hyllest. Diktet opererer som ein sofistikert intellektuell intervensjon som korrigerer Said p\u00e5 fleire avgjerande punkt.")

    # I. Identitet som skaparverk
    add_subsection_heading(doc, "I", "Identitet som skaparverk: Mot essensialismen")

    add_arabic_and_translation(doc,
        "\u0625\u0646\u0651\u064e \u0627\u0644\u0647\u0648\u064a\u0629\u064e \u0628\u0646\u062a\u064f \u0627\u0644\u0648\u0644\u0627\u062f\u0629\n\u0644\u0643\u0646\u0647\u0627 \u0641\u064a \u0627\u0644\u0646\u0647\u0627\u064a\u0629 \u0625\u0628\u062f\u0627\u0639\u064f \u0635\u0627\u062d\u0628\u0647\u0627\n\u0644\u0627 \u0648\u0631\u0627\u062b\u0629 \u0645\u0627\u0636\u064d",
        "Identitet er f\u00f8dselen si dotter, men til slutt er ho skaparverket til eigaren sin, ikkje arv fr\u00e5 fortida.")

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(3)
    p.add_run("Dette er ein ")
    run = p.add_run("anti-essensialistisk")
    run.italic = True
    p.add_run(" posisjon som direkte motseier det Chibber (2020) identifiserer som Said sitt \u00abandr argument\u00bb. Om identitet er noko som blir ")
    run = p.add_run("skapt")
    run.italic = True
    p.add_run(" av individet i m\u00f8te med materielle og historiske omstende, d\u00e5 kan det ikkje eksistere nokon \u00ablatent orientalisme\u00bb som er innebygd i vestleg kultur sidan Aiskhylos.")

    # II. Austen er ikkje heilt Aust
    add_subsection_heading(doc, "II", "\u00abAusten er ikkje heilt Aust\u00bb: Oppheving av bin\u00e6ren")

    add_arabic_and_translation(doc,
        "\u0644\u0627 \u0627\u0644\u0634\u0631\u0642\u064f \u0634\u0631\u0642\u064c \u062a\u0645\u0627\u0645\u0627\u064b \u0648\u0644\u0627 \u0627\u0644\u063a\u0631\u0628\u064f \u063a\u0631\u0628\u064c \u062a\u0645\u0627\u0645\u0627\u064b\n\u0644\u0623\u0646 \u0627\u0644\u0647\u0648\u064a\u0651\u0629\u064e \u0645\u0641\u062a\u0648\u062d\u0629\u064c \u0644\u0644\u062a\u0639\u062f\u0651\u064f\u062f\n\u0644\u0627 \u0642\u0644\u0639\u0629\u064b \u0623\u0648 \u062e\u0646\u0627\u062f\u0642\u064e",
        "Austen er ikkje heilt Aust, og Vesten er ikkje heilt Vest, for identiteten er open for mangfald: ikkje ei festning eller vollgraver.")

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(3)
    p.add_run("Al-Azm (2000) p\u00e5peikte allereie i 1981 at Said, ved \u00e5 postulere ein grunnleggjande forskjell mellom \u00abVesten\u00bb og \u00abOrienten\u00bb, risikerer \u00e5 reprodusere den same dikotomien som orientalismen sj\u00f8lv. Darwish sin Said avviser denne dikotomien fullstendig. Identitet er ")
    run = p.add_run("open")
    run.italic = True
    p.add_run(", ")
    run = p.add_run("por\u00f8s")
    run.italic = True
    p.add_run(", ")
    run = p.add_run("fleirfaldig")
    run.italic = True
    p.add_run(".")

    # III. Blodet som brot
    add_subsection_heading(doc, "III", "Blodet som brot: Det materielle si inntrengning")

    add_arabic_and_translation(doc,
        "\u062f\u0645\u064c, \u0648\u062f\u0645\u064c\u060c \u0648\u062f\u0645\u064c \u0641\u064a \u0628\u0644\u0627\u062f\u0643\n\u0641\u064a \u0627\u0633\u0645\u064a \u0648\u0641\u064a \u0627\u0633\u0645\u0643\u060c \u0641\u064a \u0632\u0647\u0631\u0629 \u0627\u0644\u0644\u0648\u0632\u060c \u0641\u064a \u0642\u0634\u0631\u0629 \u0627\u0644\u0645\u0648\u0632\n\u0641\u064a \u0644\u0628\u0646 \u0627\u0644\u0637\u0641\u0644\u060c \u0641\u064a \u0627\u0644\u0636\u0648\u0621 \u0648\u0627\u0644\u0638\u0644\u0651",
        "Blod, og blod, og blod i landet ditt: i mitt namn og i ditt namn, i mandelblomen, i bananskalet, i barnemelka, i lyset og skuggen\u2026")

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(3)
    p.add_run("Repetisjonen av \u00abblod\u00bb (\u062f\u0645) fungerer som eit formelt brot i diktet. Det vakre, det metaforiske, det kulturelle blir avbrote av r\u00e5 materialitet. Denne passasjen f\u00e5r ei smertelig aktualitet i lys av situasjonen i Gaza sidan oktober 2023. Den internasjonale domstolen (ICJ, 2024) har p\u00e5lagt Israel \u00e5 sikre humanit\u00e6r hjelp, medan Amnesty International (2024) og Human Rights Watch (2024) har dokumentert det dei kallar folkemord. Per september 2025 har over 65\u2009000 palestinarar mista livet, av desse minst 19\u2009000 born (Srinivas et al., 2025).")

    # IV. Sodom sin krig
    add_subsection_heading(doc, "IV", "\u00abSodom sin krig mot folket i Babylon\u00bb")

    add_arabic_and_translation(doc,
        "\u0643\u0627\u0646 \u064a\u0642\u0627\u0648\u0645 \u062d\u064e\u0631\u0652\u0628\u064e \u0633\u064e\u062f\u064f\u0648\u0645\u064e \u0639\u0644\u0649 \u0623\u0647\u0644 \u0628\u0627\u0628\u0644 \u0648\u0627\u0644\u0633\u0631\u0637\u0627\u0646 \u0645\u0639\u0627\u064b",
        "Han kjempa mot Sodom sin krig mot folket i Babylon og kreften samstundes.")

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(3)
    p.add_run("Amerika er Sodom; Irak er Babylon. Said d\u00f8ydde av leukemi medan han motsette seg invasjonen av Irak i 2003. Darwish nektar \u00e5 skilje det personlege fr\u00e5 det politiske, det kroppslege fr\u00e5 det geopolitiske. Denna linja f\u00e5r forsterka relevans i 2025/2026 med Trump-administrasjonen si gjenoppliving av Monroe-doktrinen og bortf\u00f8ringa av Maduro (Kupchan, 2026).")

    # V. Tenkjaren og poeten
    add_subsection_heading(doc, "V", "Tenkjaren og poeten: Eit spenningsforhold")

    add_arabic_and_translation(doc,
        "\u0627\u0644\u0645\u064f\u0641\u064e\u0643\u0651\u064e\u0631\u064f \u064a\u0643\u0628\u064e\u062d\u064f \u0633\u064e\u0631\u0652\u062f\u064e \u0627\u0644\u0631\u0648\u0627\u0626\u064a\u0651\n\u0648\u0627\u0644\u0641\u064a\u0644\u0633\u0648\u0641 \u064a\u064f\u0634\u064e\u0631\u0651\u0650\u062d\u064f \u0648\u064e\u0631\u0652\u062f\u064e \u0627\u0644\u0645\u064f\u063a\u064e\u0646\u0651\u064a",
        "Tenkjaren held attende forteljaren sitt narrativ, og filosofen dissekerer songaren si rose.")

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(3)
    p.add_run("Dette er Darwish sin milde kritikk av Said, og implisitt av akademisk diskurs som s\u00e5dan. Tenkjaren (Said som teoretikar) held attende forteljaren; filosofen dissekerer poeten si blome. Analyse drep det han analyserer.")

    # VI. Det umoglege
    add_subsection_heading(doc, "VI", "Det umoglege som politisk horisont")

    add_arabic_and_translation(doc,
        "\u0625\u0646 \u0645\u062a\u0651\u064f \u0642\u0628\u0644\u0643\n\u0623\u064f\u0648\u0635\u064a\u0643\u064e \u0628\u0627\u0644\u0645\u064f\u0633\u062a\u062d\u064a\u0644\u0652\n\u0633\u0623\u0644\u062a\u064f: \u0647\u0644 \u0627\u0644\u0645\u064f\u0633\u062a\u062d\u064a\u0644\u064f \u0628\u064e\u0639\u064a\u062f\u061f\n\u0641\u0642\u0627\u0644: \u0639\u0644\u0649 \u0628\u064f\u0639\u062f\u0650 \u062c\u064a\u0644\u0652",
        "Om eg d\u00f8yr f\u00f8r deg / Gjev eg deg det umoglege i arv / Eg spurde: Er det umoglege langt unna? / Han svara: Ein generasjon unna.")

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.space_after = Pt(3)
    p.add_run("(Darwish, 2003/2005, mi omsetjing)")
    p.paragraph_format.first_line_indent = Cm(2.5)

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(3)
    p.add_run("Det umoglege er ")
    run = p.add_run("ein generasjon unna")
    run.italic = True
    p.add_run(". Ikkje evig. Ikkje metafysisk. Ikkje innebygd i ei tidlaus kulturell disposisjon. Det umoglege er historisk situert. Det kan oppn\u00e5ast gjennom politisk kamp, ikkje gjennom diskursanalyse.")


def build_section_13(doc):
    """13. VERKET SOM ARGUMENT: PALIMPSESTEN SI DOBLE LESING"""
    add_section_heading(doc, 13, "Verket som argument: Palimpsesten si doble lesing")

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.space_after = Pt(3)
    p.add_run("Diktet ")
    run = p.add_run("er")
    run.italic = True
    p.add_run(" palimpsesten som kunstverket fors\u00f8kjer \u00e5 realisere visuelt. Det held to register samstundes: overflatelaget (kulturelt) med sj\u00f8lvets mangfald, metaforar, det estetiske som fridom, og djupnelaget (materielt) med blod, kreft, krig, snikskyttarar, monsteret av sanning, Sodom sin krig p\u00e5 Babylon.")

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(3)
    p.add_run("Som Avanessian og T\u00f6pfer (2014) formulerer det: teikningane deira \u00abdo not aim to build a representational relationship between a pictorially correct understanding and a correlative conceptual thought\u00bb (s. 12). I staden skapar dei ein \u00aboccasion to think about thinking\u00bb.")

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(3)
    p.add_run("Avanessian (2015) si accelerationist-kritikk av akademia i ")
    run = p.add_run("Overwrite")
    run.italic = True
    p.add_run(" peikar mot det same problemet: den akademiske diskurskritikken risikerer \u00e5 bli ein komfortabel posisjon som let dei faktiske maktstrukturane ligge uutfordra. L\u00f8ysinga er ikkje meir raffinert tekstanalyse, men materiell intervensjon, det Avanessian og Hennig (2015) kallar ei poetikk som ikkje berre representerer, men ")
    run = p.add_run("produserer")
    run.italic = True
    p.add_run(" tid.")

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(3)
    p.add_run("N\u00e5r plata vert halden opp mot lyset, skjer noko spesielt: Prismene i folien bryt den svarte kalligrafien. Lyset som skin gjennom er ikkje det romantiske lyset fr\u00e5 eit illuminert manuskript, men det kalde, funksjonelle lyset som folien er laga for \u00e5 prosessere. Darwish sine ord om eksil og identitet vert fysisk inskribert i det materielle substratet for vestleg bildeproduksjon. Men plasten bryr seg ikkje om orda. Ripene i plasten er like verkelege som poesien.")

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(3)
    p.add_run("Forholdet er meir urovekkjande: dei eksisterer i parallelle verkelegheiter som knapt vedkjem kvarandre. Darwish formulerer sp\u00f8rsm\u00e5let som kunstverket stiller: \u00abEr dette landet verkeleg velsigna, eller d\u00f8ypt i blod?\u00bb Snikskyttarane treff m\u00e5la sine med utmerkt kvalitet. Diskursanalysen endrar ikkje kulens bane.")

    # Figure 11: Verket som teoretisk modell
    add_figure(doc, "figur_09.png", 11,
        "Verket som teoretisk modell (materialistisk stratigrafi). Betraktaren ser gjennom tre lag: diskurs/representasjon (poesi, kalligrafi, ideologi), industriell infrastruktur (medium, teknologi, prismefolie), og materiell basis (makt, \u00f8konomi, imperialisme). Determinansen g\u00e5r nedanfr\u00e5 og opp. Kjelde: Eige illustrasjon.",
        width=Inches(4.5))


def build_section_14(doc):
    """14. KVA ST\u00c5R ATT AV SAID?"""
    add_section_heading(doc, 14, "Kva st\u00e5r att av Said?")

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.space_after = Pt(3)
    p.add_run("Denne kritikken betyr ikkje at ")
    run = p.add_run("Orientalism")
    run.italic = True
    p.add_run(" er utan verdi. Said sitt fyrste argument, at kolonialismen produserte sin legitimerande ideologi, er b\u00e5de gyldig og viktig. Orientalistiske representasjonar finst, og dei har fungert som rettferdiggjering for dominans.")

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(3)
    p.add_run("Men Said sitt andre argument, at orientalismen er s\u00e5 djupt forankra i vestleg kultur at han i seg sj\u00f8lv genererer kolonial aggresjon, m\u00e5 avvisast. Det er ikkje berre teoretisk problematisk; det er politisk demobiliserande. Den materialistiske tradisjonen tilbyr eit anna svar. Parenti (2010) oppsummerer:")

    add_block_quote(doc,
        "In a word, empires do not just pursue \u2018power for power\u2019s sake.\u2019 There are real and enormous material interests at stake, fortunes to be made many times over.")

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(3)
    p.add_run("Desse interessene kan identifiserast, analyserast og motst\u00e5ast. Ikkje gjennom diskursanalyse, men gjennom politisk organisering, \u00f8konomisk motstand og solidaritet med dei som ber dei materielle kostnadene av imperiet.")

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(3)
    p.add_run("Said visste eigentleg dette. Hans mest effektive arbeid, journalistikken hans, dei politiske intervensjonane hans for palestinarane, forlet det teoretiske apparatet fr\u00e5 ")
    run = p.add_run("Orientalism")
    run.italic = True
    p.add_run(" heilt og dokumenterte berre fakta: okkupasjonen, fordrivinga, undertrykkinga. Det var det som talde.")


def build_conclusion(doc):
    """15. KONKLUSJON"""
    add_section_heading(doc, 15, "Konklusjon")

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.space_after = Pt(3)
    p.add_run("Det umoglege, seier Darwish til Said, er ein generasjon unna. Men generasjonen m\u00e5lest ikkje i tekstar. Han m\u00e5lest i r\u00f8yrleidningar og r\u00f8rsler. I droneangrep over Gaza. I Maduro si kidnapping. I trugselen mot Gr\u00f8nland.")

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(3)
    p.add_run("Denne artikkelen har f\u00f8rt tre hovudargument. For det fyrste: Said sitt ")
    run = p.add_run("Orientalism")
    run.italic = True
    p.add_run(" inneheld ein ul\u00f8yseleg indre motseiing mellom eit materialistisk argument (orientalismen som konsekvens av kolonialismen) og eit idealistisk argument (orientalismen som \u00e5rsak til kolonialismen), der det andre argumentet verken kan forklare imperialismens historiske m\u00f8nster eller best\u00e5 den empiriske testen Parenti formulerer. For det andre: Darwish sitt ")
    run = p.add_run("Tibaq")
    run.italic = True
    p.add_run(" korrigerer implisitt denne motseiinga gjennom ein anti-essensialistisk og materialistisk poetikk som in-")

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.space_after = Pt(3)
    p.add_run("sisterer p\u00e5 at identitet er skaparverk, ikkje arv, og at det umoglege er historisk situert. For det tredje: kunstverket \u2014 palimpsesten av poesi p\u00e5 prismefolie \u2014 visualiserer denne kritikken ved \u00e5 la infrastrukturen og diskursen eksistere side om side, synlege men uforeinlege.")

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(3)
    p.add_run("Desse tre argumenta peikar samla mot ein konklusjon med konsekvensar utover den akademiske sf\u00e6ren: den postkoloniale diskurskritikken, trass sine viktige bidrag til \u00e5 synleggjere ideologiske representasjonar, risikerer \u00e5 bli ein substitutt for den materialistiske analysen som trengs for \u00e5 forst\u00e5 \u2014 og motarbeide \u2014 imperialismens faktiske mekanismar. Prismefolien fragmenterer kalligrafien, men kalligrafien endrar ikkje prismefolien. Oljeleidningane ligg der dei ligg uavhengig av kva me skriv om dei.")

    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.first_line_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(3)
    p.add_run("Det er i dette gapet mellom diskurs og infrastruktur at b\u00e5de kunstverket og artikkelen situerer seg: ikkje for \u00e5 fornekte diskursen si rolle, men for \u00e5 insistere p\u00e5 at kritikk som ikkje konfronterer basis forblir overflate.")

    # Figure 12: Røyndom
    add_figure(doc, "figur_10.png", 12,
        "R\u00f8yndom, ikkje utanfor, spr\u00e5ket: verket insisterer p\u00e5 at det materielle og det diskursive eksisterer samstundes utan \u00e5 kansellere kvarandre. Trykkpressa som metafor: spr\u00e5ket er alltid allereie materialisert. Kjelde: Eige illustrasjon.",
        width=Inches(4.5))


def build_artwork_figures(doc):
    """Add the full-page artwork figures at the end."""
    # Figure 13: Typografisk komposisjon
    add_figure(doc, "cover arabic caligraphy.png", 13,
        "Typografisk komposisjon av dei siste linjene i Tibaq: \u00abOm eg d\u00f8yr f\u00f8r deg / Gjev eg deg det umoglege i arv / Eg spurde: Er det umoglege langt unna? / Han svara: Ein generasjon unna.\u00bb Darwish sitt testamente til Said \u2014 og til oss. Kjelde: Eige illustrasjon.",
        width=Inches(5.5))

    # Figure 1: Det ferdige verket (the 3-image collage)
    add_figure(doc, "3 image collage of the palimpsest.png", 1,
        "Det ferdige verket i tre visningar. Venstre: n\u00e6rbilete gjennom prismefolien, der ljosbrytinga fragmenterer kalligrafien i spektralfargar. \u00d8vst h\u00f8gre: verket hengt opp, sett med dagslys bakfr\u00e5. Nedst h\u00f8gre: verket belyst framanfr\u00e5, der kalligrafien trer tydelegare fram mot den industrielle overflata. Kjelde: Eige fotografi.",
        width=Inches(5.5))


def build_references(doc):
    """Build the reference list in APA 7th format."""
    add_section_heading(doc, "", "REFERANSAR")
    # Remove numbering from heading
    doc.paragraphs[-1].text = ""
    p = doc.paragraphs[-1]
    run = p.add_run("REFERANSAR")
    run.bold = True

    refs = [
        # Each ref is a list of (text, is_italic) tuples
        [
            ("ABC News. (2026, 6. januar). ", False),
            ("Trump\u2019s \u2018Donroe Doctrine\u2019 seeks influence over Western Hemisphere citing old US policy", True),
            (". https://abcnews.com/Politics/trumps-donroe-doctrine-seeks-influence-western-hemisphere-citing/story?id=128926397", False),
        ],
        [
            ("Ahmad, A. (1992). Orientalism and after: Ambivalence and metropolitan location in the work of Edward Said. I ", False),
            ("In theory: Classes, nations, literatures", True),
            (" (s. 159\u2013220). Verso.", False),
        ],
        [
            ("Ahmad, A. (2000). Between orientalism and historicism. I A. L. Macfie (Red.), ", False),
            ("Orientalism: A reader", True),
            (" (s. 285\u2013297). Edinburgh University Press.", False),
        ],
        [
            ("al-Azm, S. J. (2000). Orientalism and orientalism in reverse. I A. L. Macfie (Red.), ", False),
            ("Orientalism: A reader", True),
            (" (s. 217\u2013238). New York University Press.", False),
        ],
        [
            ("Amnesty International. (2024). ", False),
            ("Israel\u2019s genocide against Palestinians in Gaza", True),
            (". https://www.amnesty.org/en/documents/mde15/8668/2024/en/", False),
        ],
        [
            ("Avanessian, A. (2015). ", False),
            ("Overwrite: Ethics of knowledge, poetics of existence", True),
            (". Sternberg Press.", False),
        ],
        [
            ("Avanessian, A., & Hennig, A. (2015). ", False),
            ("Present tense: A poetics", True),
            (". Bloomsbury Academic.", False),
        ],
        [
            ("Avanessian, A., & T\u00f6pfer, A. (2014). ", False),
            ("Speculative drawing: 2011\u20132014", True),
            (". Sternberg Press.", False),
        ],
        [
            ("Buruma, I., & Margalit, A. (2004). ", False),
            ("Occidentalism: The West in the eyes of its enemies", True),
            (". Penguin Press.", False),
        ],
        [
            ("Chibber, V. (2018). The dual legacy of Orientalism. I B. Abu-Manneh (Red.), ", False),
            ("After Said: Postcolonial literary studies in the twenty-first century", True),
            (" (s. 37\u201352). Cambridge University Press.", False),
        ],
        [
            ("Chibber, V. (2020). Orientalism and its afterlives. ", False),
            ("Catalyst", True),
            (", ", False),
            ("4", True),
            ("(3). https://catalyst-journal.com/2020/12/orientalism-and-its-afterlives", False),
        ],
        [
            ("Chomsky, N. (2011, 1. september). The responsibility of intellectuals, redux. ", False),
            ("Boston Review", True),
            (". https://www.bostonreview.net/articles/noam-chomsky-responsibility-of-intellectuals-redux/", False),
        ],
        [
            ("Chomsky, N., & Foucault, M. (2006). ", False),
            ("The Chomsky\u2013Foucault debate: On human nature", True),
            (" (F. Elders, Red.). The New Press.", False),
        ],
        [
            ("Kupchan, C. A. (2026, 7. januar). Venezuela and beyond: Trump\u2019s \u2018America First\u2019 rhetoric masks a neo-imperialist streak. ", False),
            ("Council on Foreign Relations", True),
            (". https://www.cfr.org/expert-brief/venezuela-and-beyond-trumps-america-first-rhetoric-masks-neo-imperialist-streak", False),
        ],
        [
            ("Darwish, M. (2005). Antithesis [Tibaq] (G. El-Hage, Oms.). ", False),
            ("Journal of Arabic Literature", True),
            (", ", False),
            ("36", True),
            ("(1), 50\u201356.", False),
        ],
        [
            ("Human Rights Watch. (2024, 19. desember). ", False),
            ("Israel: Starvation used as weapon of war in Gaza", True),
            (". https://www.hrw.org/news/2024/12/19/israel-starvation-used-weapon-war-gaza", False),
        ],
        [
            ("International Association of Genocide Scholars. (2025). ", False),
            ("Resolution on Gaza", True),
            (". https://genocidescholars.org/wp-content/uploads/2025/08/IAGS-Resolution-on-Gaza-FINAL.pdf", False),
        ],
        [
            ("International Court of Justice. (2024, 26. januar). ", False),
            ("Application of the Convention on the Prevention and Punishment of the Crime of Genocide in the Gaza Strip", True),
            (" (South Africa v. Israel), Provisional Measures. https://www.icj-cij.org/case/192", False),
        ],
        [
            ("Mackay, R., & Avanessian, A. (Red.). (2014). ", False),
            ("#Accelerate: The accelerationist reader", True),
            (". Urbanomic.", False),
        ],
        [
            ("Parenti, M. (1995). ", False),
            ("Against empire", True),
            (". City Lights Books.", False),
        ],
        [
            ("Parenti, M. (2010, februar). What do empires do? ", False),
            ("Voltaire Network", True),
            (". https://www.voltairenet.org/article165118.html", False),
        ],
        [
            ("PBS NewsHour. (2026, 6. januar). ", False),
            ("How the Monroe Doctrine factors into the U.S. arrest of Venezuela\u2019s Maduro", True),
            (". https://www.pbs.org/newshour/politics/how-the-monroe-doctrine-factors-into-the-u-s-arrest-of-venezuelas-maduro", False),
        ],
        [
            ("Rahman, N. (2007). Poetry of politics and mourning: Mahmoud Darwish\u2019s genre-transforming tribute to Edward W. Said. ", False),
            ("PMLA", True),
            (", ", False),
            ("122", True),
            ("(5), 1447\u20131462. https://doi.org/10.1632/pmla.2007.122.5.1447", False),
        ],
        [
            ("Reuters. (2025, 20. februar). ", False),
            ("Timeline of Trump\u2019s remarks on Palestinian displacement, Gaza takeover", True),
            (". https://www.reuters.com/world/middle-east/timeline-trumps-remarks-palestinian-displacement-gaza-takeover-2025-02-20/", False),
        ],
        [
            ("Said, E. W. (1978). ", False),
            ("Orientalism", True),
            (". Pantheon Books.", False),
        ],
        [
            ("Said, E. W. (1993). ", False),
            ("Culture and imperialism", True),
            (". Knopf.", False),
        ],
        [
            ("Smith, D. M. (2023). ", False),
            ("Endless holocausts: Mass death in the history of the United States Empire", True),
            (". Monthly Review Press.", False),
        ],
        [
            ("Srinivas, N., Vijay, D., Alakavuklar, O. N., Shymko, Y., Bohm, S., Van Laer, K., Alamgir, F., Al-Amoudi, I., Barros, M., & Mir Zulfiqar, G. (2025). Facing the facts of the Gaza genocide: Refusing complicity, organizing solidarity. ", False),
            ("Organization", True),
            (", ", False),
            ("32", True),
            ("(8), 1089\u20131102. https://doi.org/10.1177/13505084251384843", False),
        ],
        [
            ("The White House. (2025, november). ", False),
            ("National Security Strategy of the United States of America", True),
            (". https://www.whitehouse.gov/wp-content/uploads/2025/12/2025-National-Security-Strategy.pdf", False),
        ],
    ]

    for ref_parts in refs:
        p = doc.add_paragraph(style='Normal')
        p.paragraph_format.left_indent = Cm(1.27)
        p.paragraph_format.first_line_indent = Cm(-1.27)  # Hanging indent
        p.paragraph_format.space_after = Pt(3)
        for text, is_italic in ref_parts:
            run = p.add_run(text)
            if is_italic:
                run.italic = True


def set_margins(doc):
    """Set document margins."""
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)


def main():
    doc = Document()

    # Setup
    set_margins(doc)
    setup_styles(doc)

    # Build document
    build_front_matter(doc)
    build_section_1(doc)
    build_section_2(doc)
    build_section_3(doc)
    build_section_4(doc)
    build_section_5(doc)
    build_section_6(doc)
    build_section_7(doc)
    build_section_8(doc)
    build_section_9(doc)
    build_section_10(doc)
    build_section_11(doc)
    build_section_12(doc)
    build_section_13(doc)
    build_section_14(doc)
    build_conclusion(doc)
    build_artwork_figures(doc)
    build_references(doc)

    # Save
    doc.save(OUTPUT_FILE)
    file_size = os.path.getsize(OUTPUT_FILE)
    print(f"Manuscript saved to: {OUTPUT_FILE}")
    print(f"File size: {file_size / 1024:.1f} KB ({file_size / (1024*1024):.2f} MB)")
    if file_size > 8 * 1024 * 1024:
        print("WARNING: File exceeds 8 MB limit!")
    else:
        print("OK: File is within 8 MB limit.")


if __name__ == "__main__":
    main()
